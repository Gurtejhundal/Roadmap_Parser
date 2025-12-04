import streamlit as st
import db
import parser
import themes
from datetime import datetime

import time
import io
from docx import Document
from PIL import Image, ImageDraw, ImageFont

# Page Config
st.set_page_config(page_title="Roadmap Tracer", layout="wide", initial_sidebar_state="collapsed")

# Initialize DB and UI state
if 'db_initialized' not in st.session_state:
    db.init_db()
    st.session_state['db_initialized'] = True

if 'active_page' not in st.session_state:
    st.session_state['active_page'] = "Auto Generate"

if 'nav_radio' not in st.session_state:
    st.session_state['nav_radio'] = st.session_state['active_page']

THEME_OPTIONS = [
    "Cyberpunk",
    "Sunset Neon",
    "Aqua Glow",
    "Ember Glow",
    "Titanium",
    "Anti Gravity"
]
if 'theme' not in st.session_state:
    st.session_state['theme'] = THEME_OPTIONS[0]


def apply_selected_theme(theme_name: str):
    """Inject CSS for the chosen theme."""
    if theme_name == "Cyberpunk":
        st.markdown(themes.get_cyberpunk_theme(), unsafe_allow_html=True)
    elif theme_name == "Sunset Neon":
        st.markdown(themes.get_sunset_neon_theme(), unsafe_allow_html=True)
    elif theme_name == "Aqua Glow":
        st.markdown(themes.get_aqua_glow_theme(), unsafe_allow_html=True)
    elif theme_name == "Ember Glow":
        st.markdown(themes.get_ember_glow_theme(), unsafe_allow_html=True)
    elif theme_name == "Titanium":
        st.markdown(themes.get_titanium_theme(), unsafe_allow_html=True)
    elif theme_name == "Anti Gravity":
        st.markdown(themes.get_antigravity_theme(), unsafe_allow_html=True)


def import_roadmap(name, text):
    # Normalize name and block duplicates
    chosen_name = (name or "Untitled Roadmap").strip() or "Untitled Roadmap"
    existing = db.get_roadmaps()
    if any(r['name'].strip().lower() == chosen_name.lower() for r in existing):
        st.warning("A roadmap with this name already exists. Please rename it first.")
        return False

    if not text.strip():
        st.error("Please paste some roadmap text.")
        return False

    roadmap_id = db.save_roadmap(chosen_name, text)
    parsed_items = parser.parse_roadmap(text)

    timeframe_map = {}
    count_tf = 0
    for item in parsed_items:
        if item['type'] == 'timeframe':
            parent_id = None
            if item['parent_label'] and item['parent_label'] in timeframe_map:
                parent_id = timeframe_map[item['parent_label']]

            tf_id = db.save_timeframe(roadmap_id, item['label'], item['granularity'], parent_id)
            timeframe_map[item['label']] = tf_id
            count_tf += 1

    unassigned_id = None
    count_tasks = 0
    for item in parsed_items:
        if item['type'] == 'task':
            tf_label = item['timeframe_label']
            tf_id = None

            if tf_label == "Unassigned":
                if not unassigned_id:
                    unassigned_id = db.save_timeframe(roadmap_id, "Unassigned", "generic")
                tf_id = unassigned_id
            elif tf_label in timeframe_map:
                tf_id = timeframe_map[tf_label]
            else:
                if not unassigned_id:
                    unassigned_id = db.save_timeframe(roadmap_id, "Unassigned", "generic")
                tf_id = unassigned_id

            db.save_task(tf_id, item['title'])
            count_tasks += 1

    st.success(f"Imported {count_tf} timeframes and {count_tasks} tasks!")
    return True


def build_pdf(roadmap_label, tasks):
    """Generate a simple PDF of the roadmap tasks grouped by timeframe."""
    try:
        from fpdf import FPDF
    except ImportError:
        st.error("PDF generation requires the 'fpdf' package. Install with: pip install fpdf")
        return None

    # Helper to sanitize text for latin-1
    def clean(text):
        return text.encode('latin-1', 'replace').decode('latin-1')

    # Group tasks by timeframe
    grouped = {}
    for t in tasks:
        grouped.setdefault(t['timeframe_label'], []).append(t)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, clean(roadmap_label), ln=True)
    pdf.ln(4)

    pdf.set_font("Helvetica", "", 12)
    for timeframe, items in grouped.items():
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, clean(timeframe), ln=True)
        pdf.set_font("Helvetica", "", 11)
        for idx, item in enumerate(items, start=1):
            pdf.multi_cell(0, 7, f"- {clean(item['title'])}")
        pdf.ln(2)

    pdf_bytes = pdf.output(dest="S").encode("latin-1")
    return pdf_bytes

def build_txt(roadmap_name, tasks):
    output = f"Roadmap: {roadmap_name}\n\n"
    current_group = None
    for task in tasks:
        if task['timeframe_label'] != current_group:
            current_group = task['timeframe_label']
            output += f"\n[{current_group}]\n"
        status = "[x]" if task['is_done'] else "[ ]"
        output += f"{status} {task['title']}\n"
    return output.encode('utf-8')

def build_docx(roadmap_name, tasks):
    doc = Document()
    doc.add_heading(f"Roadmap: {roadmap_name}", 0)
    
    current_group = None
    for task in tasks:
        if task['timeframe_label'] != current_group:
            current_group = task['timeframe_label']
            doc.add_heading(current_group, level=2)
        
        status = " (Done)" if task['is_done'] else ""
        doc.add_paragraph(f"{task['title']}{status}", style='List Bullet')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def build_jpeg(roadmap_name, tasks):
    # Estimate height
    line_height = 30
    header_height = 60
    padding = 20
    
    # Calculate total height roughly
    num_lines = len(tasks) + len(set(t['timeframe_label'] for t in tasks)) * 2 + 2
    img_height = (num_lines * line_height) + header_height + (padding * 2)
    img_width = 800
    
    img = Image.new('RGB', (img_width, img_height), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    
    try:
        # Try to load a nicer font, otherwise default
        font = ImageFont.truetype("arial.ttf", 16)
        header_font = ImageFont.truetype("arial.ttf", 24)
    except IOError:
        font = ImageFont.load_default()
        header_font = ImageFont.load_default()
        
    y = padding
    d.text((padding, y), f"Roadmap: {roadmap_name}", fill=(0, 0, 0), font=header_font)
    y += header_height
    
    current_group = None
    for task in tasks:
        if task['timeframe_label'] != current_group:
            current_group = task['timeframe_label']
            y += line_height
            d.text((padding, y), current_group, fill=(0, 0, 150), font=font)
            y += line_height
            
        status = "[x]" if task['is_done'] else "[ ]"
        d.text((padding + 20, y), f"{status} {task['title']}", fill=(0, 0, 0), font=font)
        y += line_height
        
    buffer = io.BytesIO()
    img.save(buffer, format="JPEG")
    return buffer.getvalue()

def handle_rename(roadmap_id: int, state_key: str):
    """Rename a roadmap directly from the editable field."""
    new_name = (st.session_state.get(state_key, "") or "").strip()
    if not new_name:
        st.warning("Please enter a name.")
        return
    existing = db.get_roadmaps()
    if any(r['id'] != roadmap_id and r['name'].strip().lower() == new_name.lower() for r in existing):
        st.warning("A roadmap with this name already exists. Please choose a different name.")
        return
    db.rename_roadmap(roadmap_id, new_name)
    st.success("Roadmap renamed!")
    st.rerun()


def generate_roadmap_text(goal: str) -> str:
    """Generate a simple 4-week roadmap based on a goal string."""
    goal_clean = goal.strip() or "your goal"
    return f"""Month 1 - Foundations for {goal_clean}
Week 1 - Research and Plan
- Define the scope and success criteria for {goal_clean}
- Collect references and best practices
- Outline milestones and deliverables

Week 2 - Setup
- Set up tools, environment, and accounts needed for {goal_clean}
- Create a starter project structure
- Draft a checklist for quality and testing

Week 3 - Build Core
- Implement the core features for {goal_clean}
- Add instrumentation/logging for observability
- Run a first pass of testing and fix critical issues

Week 4 - Polish and Launch
- Refine UX/content and performance for {goal_clean}
- Add documentation/readme for how to use and extend it
- Ship a v1 and collect feedback to plan next steps
"""


def show_import_page():
    with st.form("import_form"):
        st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Name for your Roadmap</label>", unsafe_allow_html=True)
        name = st.text_input("Name", placeholder="roadmap_Name", label_visibility="collapsed")
        
        st.markdown("<div style='height: 10px'></div>", unsafe_allow_html=True)
        
        st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Paste Roadmap Text</label>", unsafe_allow_html=True)
        text = st.text_area(
            "Paste Roadmap Text",
            height=300,
            placeholder="Month 1...\nWeek 1...\n- Task 1",
            label_visibility="collapsed"
        )

        st.markdown("<div style='height: 10px'></div>", unsafe_allow_html=True)
        submitted = st.form_submit_button("Import and Analyze", type="primary", use_container_width=True)

        if submitted:
            success = import_roadmap(name, text)
            if success:
                msg_placeholder = st.empty()
                msg_placeholder.success("Analyze completed! Redirecting...")
                time.sleep(2)
                st.session_state['force_redirect'] = "View Tasks"  # Request redirect for next run
                st.rerun()


def show_auto_page():
    if 'auto_roadmap_text' not in st.session_state:
        st.session_state['auto_roadmap_text'] = ""

    with st.form("auto_form"):
        st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Name for your Roadmap</label>", unsafe_allow_html=True)
        name = st.text_input("Name", placeholder="roadmap_Name", label_visibility="collapsed")
        
        st.markdown("<div style='height: 10px'></div>", unsafe_allow_html=True)
        
        col_goal, col_btn = st.columns([0.8, 0.2], gap="small")
        with col_goal:
            st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Describe your goal</label>", unsafe_allow_html=True)
            goal_prompt = st.text_input("Goal", placeholder="give me a roadmap so failure is 100% my fault", label_visibility="collapsed")
        with col_btn:
            st.markdown("<div style='height: 29px'></div>", unsafe_allow_html=True)
            generate_clicked = st.form_submit_button("Generate", type="secondary", use_container_width=True)

        st.markdown("<div style='height: 10px'></div>", unsafe_allow_html=True)
        
        st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Generated Roadmap (editable)</label>", unsafe_allow_html=True)
        text = st.text_area(
            "Generated Roadmap",
            height=300,
            placeholder="Month 1...\nWeek 1...\n- Task 1",
            value=st.session_state.get('auto_roadmap_text', ""),
            label_visibility="collapsed"
        )

        st.markdown("<div style='height: 10px'></div>", unsafe_allow_html=True)
        submitted = st.form_submit_button("Analyze", type="primary", use_container_width=True)

        if generate_clicked:
            st.session_state['auto_roadmap_text'] = generate_roadmap_text(goal_prompt or name or "your goal")
            st.rerun()

        if submitted:
            st.session_state['auto_roadmap_text'] = text
            success = import_roadmap(name, text)
            if success:
                msg_placeholder = st.empty()
                msg_placeholder.success("Analyze completed! Redirecting...")
                time.sleep(2)
                st.session_state['force_redirect'] = "View Tasks"  # Request redirect for next run
                st.rerun()


                st.session_state['force_redirect'] = "View Tasks"  # Request redirect for next run
                st.rerun()


def show_edit_page():
    st.header("Edit Roadmap")
    
    if 'edit_roadmap_id' not in st.session_state:
        st.error("No roadmap selected for editing.")
        if st.button("Back to View"):
            st.session_state['force_redirect'] = "View Tasks"
            st.rerun()
        return

    roadmap_id = st.session_state['edit_roadmap_id']
    current_text = st.session_state.get('edit_roadmap_text', "")
    
    with st.form("edit_form"):
        new_text = st.text_area(
            "Edit Roadmap Content",
            height=400,
            value=current_text,
            help="Modify your tasks here. Headers are timeframes, lines are tasks."
        )
        
        c1, c2 = st.columns([0.2, 0.8])
        with c1:
            submitted = st.form_submit_button("Update Roadmap", type="primary")
        with c2:
            cancelled = st.form_submit_button("Cancel", type="secondary")
            
        if submitted:
            # We need to pass the parser function. 
            # Since we are in app.py and imported parser, we can pass parser.parse_roadmap
            success = db.update_roadmap_content(roadmap_id, new_text, parser.parse_roadmap)
            if success:
                st.success("Roadmap updated successfully!")
                time.sleep(1)
                st.session_state['force_redirect'] = "View Tasks"
                st.rerun()
            else:
                st.error("Failed to update roadmap.")
                
        if cancelled:
            st.session_state['force_redirect'] = "View Tasks"
            st.rerun()


def show_view_page():
    roadmaps = db.get_roadmaps()
    if not roadmaps:
        st.info("No roadmaps found. Go to 'Import Roadmap' to add one.")
        return

    def fmt_date(raw):
        try:
            return datetime.fromisoformat(raw).strftime("%d-%b-%Y").lower()
        except Exception:
            return raw[:10]

    # Header Layout
    col_head_title, col_head_prog = st.columns([0.6, 0.4])
    with col_head_title:
        st.markdown(f"<h1 class='page-title'>Your Tasks</h1>", unsafe_allow_html=True)

    # 4-Column Layout
    c1, c2, c3, c4 = st.columns(4, gap="medium")

    # Column 1: Active Project
    with c1:
        st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Active Project</label>", unsafe_allow_html=True)
        selected_roadmap = st.selectbox(
            "Active Project",
            roadmaps,
            format_func=lambda r: r['name'],
            label_visibility="collapsed"
        )
        selected_roadmap_id = selected_roadmap['id']
        selected_roadmap_name = selected_roadmap['name']
        
        st.markdown(f"<div style='margin-top: 5px; margin-bottom: 10px; color:#a0a8b8; font-size: 0.8rem;'>Created: {fmt_date(selected_roadmap['created_at'])}</div>", unsafe_allow_html=True)
        
        if st.button("Edit Roadmap", type="secondary", use_container_width=True):
             # Reconstruct text for editing
             tasks = db.get_all_tasks_for_roadmap(selected_roadmap_id)
             raw_text = selected_roadmap.get('raw_text', "")
             if not raw_text:
                 current_group = None
                 reconstructed = ""
                 for t in tasks:
                     if t['timeframe_label'] != current_group:
                         current_group = t['timeframe_label']
                         reconstructed += f"\n{current_group}\n"
                     reconstructed += f"{t['title']}\n"
                 raw_text = reconstructed.strip()
             
             st.session_state['edit_roadmap_id'] = selected_roadmap_id
             st.session_state['edit_roadmap_text'] = raw_text
             st.session_state['force_redirect'] = "Edit Roadmap"
             st.rerun()

    # Now that we have selected_roadmap_id, we can render the progress bar in the header
    with col_head_prog:
        # Progress Bar
        all_tasks_for_progress = db.get_tasks(selected_roadmap_id, None)
        total_tasks = len(all_tasks_for_progress)
        completed_tasks = sum(1 for t in all_tasks_for_progress if t['is_done'])
        progress = completed_tasks / total_tasks if total_tasks > 0 else 0
        
        st.markdown(f"<div style='text-align: right; margin-bottom: 5px; color:#e7edf7; font-size: 0.9rem; font-weight: 600;'>Progress: {int(progress*100)}%</div>", unsafe_allow_html=True)
        st.progress(progress)

    # Column 2: Timeline Filter
    with c2:
        st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Timeline Filter</label>", unsafe_allow_html=True)
        # Get timeframes without granularity filter
        timeframes = db.get_timeframes(selected_roadmap_id, None)
        timeframe_options = {"All": None}
        for tf in timeframes:
            timeframe_options[tf['label']] = tf['id']
            
        selected_tf_label = st.selectbox("Timeline Filter", list(timeframe_options.keys()), label_visibility="collapsed")
        selected_tf_id = timeframe_options[selected_tf_label]
        
        # Align with the buttons in other columns
        st.markdown("<div style='height: 28px'></div>", unsafe_allow_html=True)
        
        if st.button("Delete", type="primary", use_container_width=True):
            db.delete_roadmap(selected_roadmap_id)
            st.success("Roadmap deleted!")
            st.rerun()

    # Column 3: Download Roadmap
    with c3:
        st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Download Roadmap</label>", unsafe_allow_html=True)
        dl_fmt = st.selectbox("Format", ["PDF", "TXT", "DOCX", "JPEG"], label_visibility="collapsed")
        
        # Align with the buttons in other columns
        # Created date text is roughly 2 lines height or 30px. 
        # Let's try 28px to align with the button in Col 1 which is under the date.
        st.markdown("<div style='height: 28px'></div>", unsafe_allow_html=True)
        
        # all_tasks_for_progress is already fetched above
        has_tasks = bool(all_tasks_for_progress)
        
        if has_tasks:
            data = None
            mime = ""
            ext = ""
            
            if dl_fmt == "PDF":
                data = build_pdf(selected_roadmap_name, all_tasks_for_progress)
                mime = "application/pdf"
                ext = "pdf"
            elif dl_fmt == "TXT":
                data = build_txt(selected_roadmap_name, all_tasks_for_progress)
                mime = "text/plain"
                ext = "txt"
            elif dl_fmt == "DOCX":
                data = build_docx(selected_roadmap_name, all_tasks_for_progress)
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ext = "docx"
            elif dl_fmt == "JPEG":
                data = build_jpeg(selected_roadmap_name, all_tasks_for_progress)
                mime = "image/jpeg"
                ext = "jpg"
            
            if data:
                st.download_button(
                    "Save",
                    data,
                    file_name=f"{selected_roadmap_name}.{ext}",
                    mime=mime,
                    type="secondary",
                    use_container_width=True
                )
        else:
             st.button("Save", disabled=True, use_container_width=True)

    # Column 4: Edit Name
    with c4:
        st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Edit Name</label>", unsafe_allow_html=True)
        
        name_key = f"roadmap_name_{selected_roadmap_id}"
        st.text_input(
            "Edit name",
            value=selected_roadmap_name,
            key=name_key,
            on_change=handle_rename,
            args=(selected_roadmap_id, name_key),
            label_visibility="collapsed",
            placeholder="Rename project..."
        )

    st.divider()

    tasks = db.get_tasks(selected_roadmap_id, selected_tf_id)

    if not tasks:
        st.info("No tasks found for this selection.")
    else:
        current_group = None
        for task in tasks:
            if task['timeframe_label'] != current_group:
                current_group = task['timeframe_label']
                st.markdown("<div style='height: 20px'></div>", unsafe_allow_html=True)
                st.markdown(f"### {current_group}")

            c1, c2 = st.columns([0.05, 0.95])
            with c1:
                is_done = st.checkbox("", value=bool(task['is_done']), key=f"task_{task['id']}")
                if is_done != bool(task['is_done']):
                    db.update_task_status(task['id'], is_done)
                    st.rerun()
            with c2:
                if task['is_done']:
                    st.markdown(f"~~{task['title']}~~")
                else:
                    st.markdown(task['title'])


def main():
    # Handle forced redirection (to avoid modifying widgets after instantiation)
    if 'force_redirect' in st.session_state:
        target_page = st.session_state['force_redirect']
        st.session_state['active_page'] = target_page
        st.session_state['nav_radio'] = target_page
        del st.session_state['force_redirect']
        st.rerun()

    # UI shell styling for nav dots, theme list, and slide animation
    st.markdown(
        """
        <style>
            /* section[data-testid="stSidebar"] {display: none;} */
            .main .block-container {max-width: 1250px; padding: 0.8rem 1rem 1.8rem 1rem;}
            .page-title {margin: 0 0 1rem 0; font-size: 40px; font-weight: 800; color: #e7edf7;}
            /* tighten column padding */
            div[data-testid="column"] {padding-left: 6px !important; padding-right: 6px !important;}
            /* bolder labels */
            label[data-testid="stMetricLabel"], .stTextInput label, .stTextArea label, .stSelectbox label {
                font-size: 16px !important;
                font-weight: 800 !important;
                color: #e7edf7 !important;
            }
            /* Global Button Styling - Unified */
            button[data-testid^="stBaseButton-"] {
                border: none !important;
                background: linear-gradient(90deg, var(--primary), var(--accent)) !important;
                color: #000000 !important;
                font-weight: 700 !important;
                text-transform: uppercase;
                letter-spacing: 0.05em;
                border-radius: 8px !important;
                padding: 0.5rem 1rem !important;
                box-shadow: 0 4px 14px rgba(0,0,0,0.25);
                transition: transform 0.15s ease, box-shadow 0.2s ease;
            }
            
            button[data-testid^="stBaseButton-"]:hover {
                transform: translateY(-2px);
                /* Use theme primary color for glow */
                box-shadow: 0 0 20px var(--primary) !important; 
            }

            /* Force text styling inside buttons */
            button[data-testid^="stBaseButton-"] p, 
            button[data-testid^="stBaseButton-"] div,
            .stDownloadButton > button p,
            .stDownloadButton > button div {
                color: #000000 !important;
            }

            /* Active/Click Effect */
            button[data-testid^="stBaseButton-"]:active,
            .stDownloadButton > button:active {
                transform: translateY(1px) !important;
                box-shadow: 0 0 10px var(--primary) !important;
            }
            /* keep download button simple - actually unify it too */
            .stDownloadButton>button {
                border: none !important;
                background: linear-gradient(90deg, var(--primary), var(--accent)) !important;
                color: #000000 !important;
                font-weight: 700 !important;
                text-transform: uppercase;
                letter-spacing: 0.05em;
                border-radius: 8px !important;
                padding: 0.5rem 1rem !important;
                box-shadow: 0 4px 14px rgba(0,0,0,0.25);
                transition: transform 0.15s ease, box-shadow 0.2s ease;
            }
            .stDownloadButton>button:hover {
                transform: translateY(-2px);
                box-shadow: 0 0 20px var(--primary) !important;
            }

            /* pointer cursor on selects */
            .stSelectbox label {cursor: pointer;}
            .stSelectbox div[data-baseweb="select"] > div {cursor: pointer;}

            /* Nav dots */
            .nav-dots [role="radiogroup"] {gap: 2rem !important; justify-content: center;}
            .nav-dots label {
                padding: 6px 10px 6px 28px;
                margin: 0;
                position: relative;
                cursor: pointer;
                font-size: 22px;
                font-weight: 700;
                letter-spacing: 0.01em;
                color: #e7edf7;
            }
            .nav-dots label div:first-of-type {display: none;} /* hide default radio icon */
            .nav-dots label div:last-of-type {
                color: inherit;
                font-size: inherit;
                font-weight: inherit;
            }
            .nav-dots label::before {
                content: "";
                position: absolute;
                left: 0;
                top: 50%;
                transform: translateY(-50%);
                width: 14px;
                height: 14px;
                border-radius: 50%;
                border: 2px solid #6b768a;
                background: transparent;
            }
            .nav-dots label:has(input:checked)::before {
                background: var(--primary);
                border-color: var(--primary);
                box-shadow: 0 0 0 4px rgba(255, 255, 255, 0.1);
            }



            .page-shell {animation: slideIn 0.35s ease;}
            @keyframes slideIn {from {opacity: 0; transform: translateY(8px);} to {opacity: 1; transform: translateY(0);}}
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("<h1 style='text-align: center; margin-bottom: 2rem;'>Roadmap Tracer</h1>", unsafe_allow_html=True)

    apply_selected_theme(st.session_state['theme'])

    # Navigation and Theme Selector
    col_nav, col_theme = st.columns([0.75, 0.25])
    
    with col_nav:
        st.markdown('<div class="nav-dots">', unsafe_allow_html=True)
        nav_pages = ["Auto Generate", "Import Roadmap", "View Tasks"]
        if st.session_state.get('active_page') == "Edit Roadmap":
            nav_pages.append("Edit Roadmap")
        nav_choice = st.radio(
            "Navigation",
            nav_pages,
            horizontal=True,
            label_visibility="collapsed",
            key="nav_radio",
        )
        st.markdown('</div>', unsafe_allow_html=True)
        st.session_state['active_page'] = nav_choice

    with col_theme:
        st.markdown("<label style='font-size: 16px; font-weight: 800; color: #e7edf7; display: block; margin-bottom: 0.5rem;'>Themes</label>", unsafe_allow_html=True)
        theme_choice = st.selectbox(
            "Theme",
            THEME_OPTIONS,
            index=THEME_OPTIONS.index(st.session_state['theme']),
            label_visibility="collapsed"
        )
        if theme_choice != st.session_state['theme']:
            st.session_state['theme'] = theme_choice
            st.rerun()

    if st.session_state['active_page'] == "Import Roadmap":
        st.markdown(f"<h1 class='page-title'>Roadmap</h1>", unsafe_allow_html=True)
    elif st.session_state['active_page'] == "Auto Generate":
        st.markdown(f"<h1 class='page-title'>Auto Generate</h1>", unsafe_allow_html=True)
    # View Tasks handles its own title now

    # Theme selector removed from sidebar
    # with st.sidebar:
    #     st.title("Settings")
    #     theme_choice = st.selectbox(
    #         "Choose Theme",
    #         THEME_OPTIONS,
    #         index=THEME_OPTIONS.index(st.session_state['theme']),
    #     )
    #     if theme_choice != st.session_state['theme']:
    #         st.session_state['theme'] = theme_choice
    #         st.rerun()

    st.markdown('<div class="page-shell">', unsafe_allow_html=True)
    if st.session_state['active_page'] == "Import Roadmap":
        show_import_page()
    elif st.session_state['active_page'] == "Auto Generate":
        show_auto_page()
    elif st.session_state['active_page'] == "View Tasks":
        show_view_page()
    elif st.session_state['active_page'] == "Edit Roadmap":
        show_edit_page()
    st.markdown('</div>', unsafe_allow_html=True)


if __name__ == "__main__":
    main()
